from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"D:\kimi\agribusiness_finalized\docs\AgriBusiness-Platform-Audit-August-2026.docx"

def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tc_pr.append(shd)

def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None: tc_w = OxmlElement('w:tcW'); tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(width)); tc_w.set(qn('w:type'), 'dxa')

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.style = 'Table Grid'
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = head; shade(cell, 'E8EEF5'); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs: run.bold = True
        if widths: set_cell_width(cell, widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value; cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths: set_cell_width(cells[i], widths[i])
    doc.add_paragraph()
    return t

def bullets(doc, items):
    for item in items: doc.add_paragraph(item, style='List Bullet')

doc = Document()
sec = doc.sections[0]; sec.top_margin = Inches(1); sec.bottom_margin = Inches(1); sec.left_margin = Inches(1); sec.right_margin = Inches(1)
styles = doc.styles
styles['Normal'].font.name = 'Calibri'; styles['Normal'].font.size = Pt(11)
styles['Normal'].paragraph_format.space_after = Pt(6); styles['Normal'].paragraph_format.line_spacing = 1.1
for name, size, color in [('Heading 1',16,'2E74B5'),('Heading 2',13,'2E74B5'),('Heading 3',12,'1F4D78')]:
    style = styles[name]; style.font.name = 'Calibri'; style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(12); style.paragraph_format.space_after = Pt(6)

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run('AgriBusiness Platform Audit'); r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(11,37,69)
sub = doc.add_paragraph('Product, frontend, backend, security and operations review | 24 August 2026')
sub.runs[0].font.color.rgb = RGBColor(85,85,85)

doc.add_heading('Executive summary', 1)
doc.add_paragraph('AgriBusiness is a strong Pakistan-focused marketplace foundation: it already contains account roles, profiles, category browsing, listings, project briefs, proposals, messaging, clinic-style support, internal advertising, approval controls, trial handling, and administrative governance. The product is not yet release-complete because several critical database-to-UI mismatches, incomplete monetization workflows, and trust/safety gaps can lead to blank screens, misleading metrics, or unverified business activity.')
table(doc, ['Area','Assessment','Priority'], [
 ['Core marketplace','Implemented foundation: listings, profiles, categories, messaging, projects and proposals.','Good base'],
 ['Role onboarding','Now captures sector, specialities and role-specific information. Migration must be deployed.','Release blocker until deployed'],
 ['Admin operations','Controls exist, but data sources had table-name mismatches and lacked partial-failure feedback.','Fixed in code; verify in staging'],
 ['Payments / project readiness','Projects accept proposals but do not hold funds, milestones, payouts or disputes.','High'],
 ['Trust & moderation','RLS and audit model exist; verification evidence and enforcement workflows need completion.','High'],
], [2300, 5200, 1860])

doc.add_heading('What the website currently does', 1)
bullets(doc, [
 'Connects farmers/producers, buyers/traders, consultants, companies and students through role-specific workspaces.',
 'Supports directory discovery, keyword matching, commercial listings, buying requirements, open project briefs, consultant proposals and consented direct-contact requests.',
 'Provides crop/animal clinic surfaces, market-rate pages, community feed functionality, notifications and direct messages.',
 'Uses a category hierarchy for products and services, with ad categories and geographic targeting.',
 'Provides internal advertising plans, pending/approved/rejected campaign states, manual review and scheduled weekly rotation.',
 'Provides a seven-day account trial in database logic, with expiry automation and subscription/payment data models.',
])

doc.add_heading('Issues corrected in this review', 1)
table(doc, ['Issue','Impact','Correction'], [
 ['Admin queried admin_audit_logs','Audit tab would return an error or remain empty because the schema table is admin_audit_log.','Updated query to the schema name and added a visible partial-load alert.'],
 ['Admin/profile queried rfps','Project telemetry, project deletion and profile project portfolios used a non-existent table.','Updated queries and deletion action to use projects.'],
 ['Admin chart used invented history','A live cockpit displayed fabricated monthly values, damaging decision-making trust.','Replaced with actual seven-day new-member and new-listing counts based on created_at.'],
 ['Profile saves ignored database errors','The UI could announce success after a failed profile/private-contact write.','Profile and private-record write errors now stop the success state.'],
 ['Signup created empty role records','New users had limited discoverability and companies had no organization record.','Role-aware onboarding migration persists details and initializes company organizations.'],
], [2300, 2700, 4360])

doc.add_heading('Portal-by-portal review', 1)
doc.add_heading('Admin portal', 2)
bullets(doc, [
 'Strengths: clear role gate, separated admin login, content/ad/member/rates surfaces, audit model and moderation RPCs.',
 'Improve: replace client-side aggregate counts with administrator RPCs or analytics views; current list limits can undercount platform totals.',
 'Improve: show query-specific errors and a last-successful sync time on every tab, not only a general warning.',
 'Improve: require moderation reason, reviewer identity and before/after state for every destructive or approval action; add undo where feasible.',
 'Improve: make charts explicitly label their time window and show an empty state rather than a zero-looking chart when no events exist.',
])
doc.add_heading('Member dashboard', 2)
bullets(doc, [
 'Strengths: genuine role-specific forms and workspace language; server-side role policies prevent obvious cross-role publishing.',
 'Improve: direct dashboard quick actions to the exact active workspace tab or dedicated composer route; generic /dashboard links are ambiguous.',
 'Improve: add draft autosave, image upload/reorder, field-level validation, progress indicators and clear publication status.',
 'Improve: use a dedicated portfolio/gallery model rather than treating a single URL or listing images as a portfolio.',
 'Improve: generate recommendations only after profile completeness reaches a useful threshold, and explain why each match appears.',
])
doc.add_heading('Public profile', 2)
bullets(doc, [
 'Strengths: private contact is protected behind accepted connections, and public directory data is intentionally limited.',
 'Improve: show primary sector, verified credentials, services/products, response availability and portfolio work in a structured public layout.',
 'Improve: give owners an explicit save error/success state for each subsection and avoid presenting unpublished/private fields publicly.',
 'Improve: add report/block controls, profile availability status and a structured verification request workflow.',
])

doc.add_heading('Backend and security assessment', 1)
table(doc, ['Topic','Current state','Recommendation'], [
 ['Authorization','RLS, role checks, private-profile separation and admin RPCs are present.','Run policy tests against every table/RPC; use service-role only in protected Edge Functions.'],
 ['Schema consistency','Migrations are rich, but UI table names drifted from schema names.','Generate typed Supabase types and add CI checks for every queried relation.'],
 ['Search','Keyword and semantic-search components exist.','Run embedding generation on profile/listing change; provide lexical fallback and explain matches.'],
 ['Ads','Approval, rotation and targeting structures exist.','Enforce plan limits, impression caps, ad expiry and advertiser eligibility server-side.'],
 ['Payments','Subscription/payment records and webhooks exist.','Use verified webhook signatures, idempotency keys, reconciliation jobs and immutable payment events.'],
 ['Project payments','No escrow/milestone workflow.','Add project contract, funded escrow, milestones, acceptance, payout, refund and dispute entities before claiming ready-to-pay.'],
], [1800, 3600, 3960])

doc.add_heading('Recommended roadmap', 1)
table(doc, ['When','Outcome','Work'], [
 ['Before production','Trustworthy core','Deploy Migration 14; test admin/profile/dashboard with real role accounts; generate Supabase types; remove fake data; add error monitoring.'],
 ['Next 2-4 weeks','Useful marketplace','Dedicated gallery/portfolio, verified-business evidence upload, saved searches, profile completeness and improved mobile testing.'],
 ['Next 1-2 months','Revenue-ready platform','Ad checkout lifecycle, plan entitlements, campaign analytics, invoice/payment reconciliation and dispute-safe project contracts.'],
 ['Later','Scale and intelligence','Semantic ranking feedback loops, notification preferences, analytics warehouse, fraud controls, moderation queues and performance budgets.'],
], [1600, 2500, 5260])

doc.add_heading('Release checklist', 1)
bullets(doc, [
 'Apply and verify all Supabase migrations in a staging project, including the new role-aware onboarding migration.',
 'Create test accounts for all five roles plus admin and exercise signup, profile edit, listing/project creation, proposal, connection, ad moderation and sign-out.',
 'Verify the database schema through automated tests; do not rely on UI screens alone.',
 'Configure Edge Function secrets, webhook signature validation, cron authentication and production error monitoring.',
 'Add a privacy policy, terms, reporting process, data retention policy and an administrator incident-response procedure.',
 'Do not market project work as “ready to pay” until an auditable payment, acceptance and payout flow exists.',
])

doc.add_heading('Design direction', 1)
doc.add_paragraph('Keep the existing field-led green palette, but make the interface feel more operational than decorative: use status, ownership, deadlines, location and verification as the primary visual hierarchy. The memorable signature should be a clear “field-to-market” activity trail that shows what a member did, what happens next and who is responsible. Use the same language for the action, confirmation and status everywhere: Publish listing -> Listing published -> Active listing.')

doc.add_paragraph('Prepared from the current source tree and route-level checks. This is an implementation audit, not a legal, penetration-testing or payment-compliance certification.').runs[0].italic = True
doc.save(OUT)
print(OUT)
